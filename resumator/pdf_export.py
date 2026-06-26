from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from html import escape as html_escape
from html.parser import HTMLParser
from typing import Iterable
import json
import os
from pathlib import Path
import re
import textwrap
import xml.etree.ElementTree as ET
from xml.sax.saxutils import escape as xml_escape
import zipfile

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas
except ImportError:  # pragma: no cover - fallback is used without reportlab
    colors = None
    letter = None
    ParagraphStyle = None
    getSampleStyleSheet = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None
    Table = None
    TableStyle = None
    pdfmetrics = None
    TTFont = None
    canvas = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError:  # pragma: no cover - optional DOCX export
    Document = None
    Inches = None
    Pt = None


PAGE_WIDTH = 612
PAGE_HEIGHT = 792
LEFT = 54
TOP = 735
BOTTOM = 58
LINE_HEIGHT = 14
BODY_FONT_SIZE = 10
APP_NAME = "Resumator 11.0"
DEVELOPER = "LEONARDO CARDOSO DE MELO TEIXEIRA MENDES - PROCURADOR FEDERAL / AGU"
DOCUMENT_TITLE = "RESUMO GERADO POR IA"
PROMPT_DOCUMENT_TITLE = "PROMPT PARA ENVIO À IA"
DOCX_IMPORT_DOCUMENT_TITLE = "DOCUMENTO DOCX IMPORTADO"


@dataclass
class RichRun:
    text: str
    bold: bool = False
    italic: bool = False
    underline: bool = False
    code: bool = False


@dataclass
class RichBlock:
    kind: str
    runs: list[RichRun] = field(default_factory=list)
    level: int = 0
    ordered: bool = False
    rows: list[list[list[RichRun]]] = field(default_factory=list)


class _ResponseHtmlParser(HTMLParser):
    BLOCK_TAGS = {"p", "div", "section", "article", "main", "blockquote"}
    STYLE_TAGS = {
        "b": "bold",
        "strong": "bold",
        "i": "italic",
        "em": "italic",
        "u": "underline",
        "code": "code",
    }

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[RichBlock] = []
        self.current_runs: list[RichRun] = []
        self.current_kind = "paragraph"
        self.current_level = 0
        self.current_ordered = False
        self.style_stack: list[dict[str, bool]] = [
            {"bold": False, "italic": False, "underline": False, "code": False}
        ]
        self.list_stack: list[str] = []
        self.table_rows: list[list[list[RichRun]]] | None = None
        self.current_row: list[list[RichRun]] | None = None
        self.current_cell: list[RichRun] | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in self.STYLE_TAGS:
            self._push_style(**{self.STYLE_TAGS[tag]: True})
            return
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self._finish_block()
            self.current_kind = "heading"
            self.current_level = int(tag[1])
            return
        if tag in self.BLOCK_TAGS:
            self._finish_block()
            self.current_kind = "paragraph"
            self.current_level = 0
            return
        if tag == "br":
            self._append_text("\n")
            return
        if tag == "pre":
            self._finish_block()
            self.current_kind = "code"
            self._push_style(code=True)
            return
        if tag in {"ul", "ol"}:
            self.list_stack.append(tag)
            return
        if tag == "li":
            self._finish_block()
            self.current_kind = "list"
            self.current_ordered = bool(self.list_stack and self.list_stack[-1] == "ol")
            self.current_level = max(len(self.list_stack) - 1, 0)
            return
        if tag == "table":
            self._finish_block()
            self.table_rows = []
            return
        if tag == "tr" and self.table_rows is not None:
            self.current_row = []
            return
        if tag in {"td", "th"} and self.current_row is not None:
            self.current_cell = []
            if tag == "th":
                self._push_style(bold=True)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.STYLE_TAGS:
            self._pop_style()
            return
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6", "p", "div", "section", "article", "main", "blockquote"}:
            self._finish_block()
            return
        if tag == "pre":
            self._finish_block()
            self._pop_style()
            return
        if tag == "li":
            self._finish_block()
            return
        if tag in {"ul", "ol"}:
            if self.list_stack:
                self.list_stack.pop()
            return
        if tag in {"td", "th"} and self.current_cell is not None and self.current_row is not None:
            if tag == "th":
                self._pop_style()
            self.current_row.append(self._normalized_runs(self.current_cell))
            self.current_cell = None
            return
        if tag == "tr" and self.current_row is not None and self.table_rows is not None:
            if self.current_row:
                self.table_rows.append(self.current_row)
            self.current_row = None
            return
        if tag == "table" and self.table_rows is not None:
            if self.table_rows:
                self.blocks.append(RichBlock(kind="table", rows=self.table_rows))
            self.table_rows = None

    def handle_data(self, data: str) -> None:
        self._append_text(data)

    def close(self) -> None:
        super().close()
        self._finish_block()

    def _push_style(self, **changes: bool) -> None:
        style = dict(self.style_stack[-1])
        style.update(changes)
        self.style_stack.append(style)

    def _pop_style(self) -> None:
        if len(self.style_stack) > 1:
            self.style_stack.pop()

    def _append_text(self, text: str) -> None:
        if not text:
            return
        style = self.style_stack[-1]
        run = RichRun(
            text=text,
            bold=style["bold"],
            italic=style["italic"],
            underline=style["underline"],
            code=style["code"],
        )
        if self.current_cell is not None:
            self.current_cell.append(run)
            return
        self.current_runs.append(run)

    def _finish_block(self) -> None:
        runs = self._normalized_runs(self.current_runs)
        if runs:
            self.blocks.append(
                RichBlock(
                    kind=self.current_kind,
                    runs=runs,
                    level=self.current_level,
                    ordered=self.current_ordered,
                )
            )
        self.current_runs = []
        self.current_kind = "paragraph"
        self.current_level = 0
        self.current_ordered = False

    @staticmethod
    def _normalized_runs(runs: list[RichRun]) -> list[RichRun]:
        merged: list[RichRun] = []
        for run in runs:
            text = run.text.replace("\r\n", "\n").replace("\r", "\n")
            if not text:
                continue
            if not run.code:
                text = re.sub(r"[ \t\f\v]+", " ", text)
            if not text.strip() and "\n" not in text:
                text = " "
            if (
                merged
                and merged[-1].bold == run.bold
                and merged[-1].italic == run.italic
                and merged[-1].underline == run.underline
                and merged[-1].code == run.code
            ):
                merged[-1].text += text
            else:
                merged.append(
                    RichRun(text, bold=run.bold, italic=run.italic, underline=run.underline, code=run.code)
                )
        while merged and not merged[0].text.strip():
            merged.pop(0)
        while merged and not merged[-1].text.strip():
            merged.pop()
        return merged


def _rich_blocks_from_html(formatted_html: str) -> list[RichBlock]:
    parser = _ResponseHtmlParser()
    try:
        parser.feed(formatted_html)
        parser.close()
    except Exception:
        return []
    return parser.blocks


def export_response_pdf(
    output_path: Path,
    response_text: str,
    prompt_name: str | None = None,
    source_pdf: Path | Iterable[Path] | None = None,
    formatted_html: str | None = None,
) -> Path:
    if formatted_html and _can_export_formatted_pdf():
        blocks = _rich_blocks_from_html(formatted_html)
        if blocks:
            _export_formatted_pdf(output_path, response_text, prompt_name, source_pdf, blocks)
            return output_path

    lines = _prepare_lines(response_text, prompt_name, source_pdf)
    if canvas is not None:
        _export_with_reportlab(output_path, lines)
        return output_path

    pages = _paginate(lines)
    pdf_bytes = _build_pdf(pages)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_bytes(pdf_bytes)
    return output_path


def export_response_docx(
    output_path: Path,
    response_text: str,
    prompt_name: str | None = None,
    source_pdf: Path | Iterable[Path] | None = None,
    document_title: str = DOCUMENT_TITLE,
    formatted_html: str | None = None,
) -> Path:
    if Document is None or Inches is None or Pt is None:
        _export_docx_fallback(output_path, response_text, prompt_name, source_pdf, document_title)
        return output_path

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    title = document.add_paragraph()
    title_run = title.add_run(document_title)
    title_run.bold = True
    title_run.font.size = Pt(14)

    for line in _metadata_lines(prompt_name, source_pdf):
        paragraph = document.add_paragraph()
        paragraph.add_run(line).italic = True

    document.add_paragraph("")

    blocks = _rich_blocks_from_html(formatted_html or "") if formatted_html else []
    if blocks:
        _add_rich_blocks_to_docx(document, blocks)
    else:
        text = response_text.strip() or "Sem resposta informada."
        for paragraph_text in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
            document.add_paragraph(paragraph_text.strip() if paragraph_text.strip() else "")

    document.save(str(output_path))
    return output_path


def export_prompt_docx(
    output_path: Path,
    prompt_text: str,
    prompt_name: str | None = None,
    source_pdf: Path | Iterable[Path] | None = None,
) -> Path:
    return export_response_docx(
        output_path,
        prompt_text,
        prompt_name=prompt_name,
        source_pdf=source_pdf,
        document_title=PROMPT_DOCUMENT_TITLE,
    )


def _add_rich_blocks_to_docx(document, blocks: list[RichBlock]) -> None:
    for block in blocks:
        if block.kind == "table":
            _add_docx_table(document, block)
            continue
        if block.kind == "heading":
            paragraph = document.add_heading(level=min(max(block.level, 1), 4))
        elif block.kind == "list":
            style = "List Number" if block.ordered else "List Bullet"
            try:
                paragraph = document.add_paragraph(style=style)
            except KeyError:
                paragraph = document.add_paragraph("1. " if block.ordered else "- ")
        else:
            paragraph = document.add_paragraph()
            if block.kind == "code":
                paragraph.paragraph_format.left_indent = Inches(0.25)
        _add_docx_runs(paragraph, block.runs)


def _add_docx_runs(paragraph, runs: list[RichRun]) -> None:
    for rich_run in runs:
        parts = rich_run.text.split("\n")
        for index, part in enumerate(parts):
            if index:
                paragraph.add_run().add_break()
            if not part:
                continue
            run = paragraph.add_run(part)
            run.bold = rich_run.bold
            run.italic = rich_run.italic
            run.underline = rich_run.underline
            if rich_run.code:
                run.font.name = "Consolas"


def _add_docx_table(document, block: RichBlock) -> None:
    rows = [row for row in block.rows if row]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            if column_index < len(row):
                _add_docx_runs(paragraph, row[column_index])
    document.add_paragraph("")


def _can_export_formatted_pdf() -> bool:
    return all(
        item is not None
        for item in (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, getSampleStyleSheet, colors, letter)
    )


def _export_formatted_pdf(
    output_path: Path,
    response_text: str,
    prompt_name: str | None,
    source_pdf: Path | Iterable[Path] | None,
    blocks: list[RichBlock],
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    regular_font, _ = _reportlab_fonts()
    normal = ParagraphStyle("ResumatorNormal", parent=styles["Normal"], fontName=regular_font, fontSize=10, leading=13)
    meta = ParagraphStyle("ResumatorMeta", parent=normal, fontSize=8.5, leading=11, italic=True)
    title_style = ParagraphStyle(
        "ResumatorTitle",
        parent=styles["Title"],
        fontName=regular_font,
        fontSize=14,
        leading=18,
        spaceAfter=8,
    )
    heading = ParagraphStyle(
        "ResumatorHeading",
        parent=styles["Heading2"],
        fontName=regular_font,
        fontSize=12,
        leading=15,
        spaceBefore=8,
        spaceAfter=4,
    )
    code_style = ParagraphStyle(
        "ResumatorCode",
        parent=normal,
        fontName="Courier",
        leftIndent=18,
        backColor=colors.whitesmoke,
    )

    elements = [Paragraph(html_escape(DOCUMENT_TITLE), title_style)]
    for line in _metadata_lines(prompt_name, source_pdf):
        elements.append(Paragraph(html_escape(line), meta))
    elements.append(Spacer(1, 10))

    if blocks:
        ordered_index = 1
        for block in blocks:
            if block.kind == "table":
                _append_pdf_table(elements, block, normal)
                ordered_index = 1
                continue
            style = heading if block.kind == "heading" else code_style if block.kind == "code" else normal
            markup = _runs_to_reportlab_markup(block.runs)
            if not markup:
                continue
            if block.kind == "list":
                prefix = f"{ordered_index}. " if block.ordered else "• "
                markup = html_escape(prefix) + markup
                ordered_index = ordered_index + 1 if block.ordered else 1
            else:
                ordered_index = 1
            elements.append(Paragraph(markup, style))
            elements.append(Spacer(1, 4))
    else:
        for paragraph in (response_text.strip() or "Sem resposta informada.").splitlines():
            elements.append(Paragraph(html_escape(paragraph.strip()), normal))
            elements.append(Spacer(1, 4))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=LEFT,
        rightMargin=LEFT,
        topMargin=58,
        bottomMargin=BOTTOM,
        title=DOCUMENT_TITLE,
    )
    doc.build(elements)


def _append_pdf_table(elements: list, block: RichBlock, normal_style) -> None:
    rows = [row for row in block.rows if row]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    data = []
    for row in rows:
        data_row = []
        for column_index in range(column_count):
            runs = row[column_index] if column_index < len(row) else []
            data_row.append(Paragraph(_runs_to_reportlab_markup(runs) or " ", normal_style))
        data.append(data_row)
    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    elements.append(table)
    elements.append(Spacer(1, 8))


def _runs_to_reportlab_markup(runs: list[RichRun]) -> str:
    parts: list[str] = []
    for run in runs:
        text = html_escape(run.text).replace("\n", "<br/>")
        if run.code:
            text = f'<font name="Courier">{text}</font>'
        if run.underline:
            text = f"<u>{text}</u>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.bold:
            text = f"<b>{text}</b>"
        parts.append(text)
    return "".join(parts).strip()


def export_docx_file_to_pdf(source_path: Path, output_path: Path) -> Path:
    source_path = Path(source_path)
    output_path = Path(output_path)
    if not source_path.exists():
        raise FileNotFoundError(f"DOCX não encontrado: {source_path}")
    if source_path.suffix.lower() != ".docx":
        raise ValueError("O arquivo selecionado não está no formato .docx.")

    word_error: Exception | None = None
    if os.name == "nt":
        try:
            if _export_docx_with_word(source_path, output_path):
                return output_path
        except Exception as exc:  # noqa: BLE001 - fallback preserves the feature without Word
            word_error = exc

    try:
        text = _extract_docx_text(source_path)
        if not text.strip():
            text = f"Documento DOCX importado: {source_path.name}"
        return export_response_pdf(
            output_path,
            text,
            prompt_name=source_path.name,
            source_pdf=None,
        )
    except Exception as exc:  # noqa: BLE001 - include Word failure when both paths fail
        if word_error is not None:
            raise RuntimeError(
                f"Falha na conversão pelo Word ({word_error}) e no fallback por texto ({exc})."
            ) from exc
        raise


def _export_docx_with_word(source_path: Path, output_path: Path) -> bool:
    try:
        import win32com.client  # type: ignore
    except Exception:
        return False

    output_path.parent.mkdir(parents=True, exist_ok=True)
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(
            str(source_path.resolve()),
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        document.SaveAs(str(output_path.resolve()), FileFormat=17)
        return output_path.exists()
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()


def _extract_docx_text(source_path: Path) -> str:
    if Document is not None:
        document = Document(str(source_path))
        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cell for cell in cells if cell))
        return "\n".join(parts)
    return _extract_docx_text_from_zip(source_path)


def _extract_docx_text_from_zip(source_path: Path) -> str:
    with zipfile.ZipFile(source_path) as archive:
        xml_data = archive.read("word/document.xml")
    root = ET.fromstring(xml_data)
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    paragraphs: list[str] = []
    for paragraph in root.findall(".//w:p", namespace):
        texts = [
            node.text or ""
            for node in paragraph.findall(".//w:t", namespace)
        ]
        if texts:
            paragraphs.append("".join(texts))
    return "\n".join(paragraphs)


def _export_docx_fallback(
    output_path: Path,
    response_text: str,
    prompt_name: str | None,
    source_pdf: Path | Iterable[Path] | None,
    document_title: str,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    paragraphs = [("title", document_title)]
    paragraphs.extend(("meta", line) for line in _metadata_lines(prompt_name, source_pdf))
    paragraphs.append(("body", ""))

    text = response_text.strip() or "Sem resposta informada."
    for paragraph_text in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        paragraphs.append(("body", paragraph_text.strip() if paragraph_text.strip() else ""))

    document_xml = _docx_document_xml(paragraphs)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", _docx_content_types_xml())
        archive.writestr("_rels/.rels", _docx_package_rels_xml())
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/_rels/document.xml.rels", _docx_document_rels_xml())


def _docx_content_types_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""


def _docx_package_rels_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""


def _docx_document_rels_xml() -> str:
    return """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>"""


def _docx_document_xml(paragraphs: list[tuple[str, str]]) -> str:
    body = "\n".join(_docx_paragraph(kind, text) for kind, text in paragraphs)
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
{body}
    <w:sectPr>
      <w:pgSz w:w="12240" w:h="15840"/>
      <w:pgMar w:top="1008" w:right="1152" w:bottom="1008" w:left="1152" w:header="720" w:footer="720" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>"""


def _docx_paragraph(kind: str, text: str) -> str:
    run_props = ""
    if kind == "title":
        run_props = "<w:rPr><w:b/><w:sz w:val=\"28\"/></w:rPr>"
    elif kind == "meta":
        run_props = "<w:rPr><w:i/><w:sz w:val=\"20\"/></w:rPr>"
    else:
        run_props = "<w:rPr><w:sz w:val=\"21\"/></w:rPr>"

    safe_text = xml_escape(_valid_xml_text(text))
    return f"""    <w:p>
      <w:r>{run_props}<w:t xml:space="preserve">{safe_text}</w:t></w:r>
    </w:p>"""


def _valid_xml_text(text: str) -> str:
    return "".join(
        char
        for char in text
        if char in "\t\n\r" or ord(char) >= 0x20
    )


def export_response_json(
    output_path: Path,
    response_text: str,
    prompt_name: str | None = None,
    source_pdf: Path | Iterable[Path] | None = None,
) -> Path:
    text = response_text.strip() or "Sem resposta informada."
    source_pdfs = _normalize_source_pdfs(source_pdf)
    payload = {
        "version": 1,
        "source": APP_NAME,
        "type": "resumo_peticao_inicial",
        "exported_at": datetime.now().isoformat(timespec="seconds"),
        "prompt": prompt_name,
        "arquivos_analisados": [path.name for path in source_pdfs],
        "caminhos_arquivos": [str(path) for path in source_pdfs],
        "resumo": text,
        "resumo_peticao": text,
        "resumo_da_peticao": text,
        "content": text,
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return output_path


def _export_with_reportlab(output_path: Path, lines: list[str]) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    regular_font, bold_font = _reportlab_fonts()
    pages = _paginate(lines)
    doc = canvas.Canvas(str(output_path), pagesize=letter)
    width, height = letter

    for page_number, page_lines in enumerate(pages, start=1):
        doc.setFont(bold_font, 13)
        doc.drawString(LEFT, height - 30, APP_NAME)
        doc.setFont(regular_font, 9)
        doc.drawString(LEFT, height - 48, f"Página {page_number} de {len(pages)}")

        y = height - 85
        for line in page_lines:
            if line == DOCUMENT_TITLE:
                doc.setFont(bold_font, 13)
            else:
                doc.setFont(regular_font, BODY_FONT_SIZE)
            doc.drawString(LEFT, y, line)
            y -= LINE_HEIGHT
        doc.showPage()

    doc.save()


def _reportlab_fonts() -> tuple[str, str]:
    if pdfmetrics is None or TTFont is None:
        return "Helvetica", "Helvetica-Bold"

    windows_dir = Path(os.environ.get("WINDIR", r"C:\Windows"))
    regular_path = windows_dir / "Fonts" / "arial.ttf"
    bold_path = windows_dir / "Fonts" / "arialbd.ttf"
    try:
        if regular_path.exists() and bold_path.exists():
            pdfmetrics.registerFont(TTFont("ResumatorArial", str(regular_path)))
            pdfmetrics.registerFont(TTFont("ResumatorArialBold", str(bold_path)))
            return "ResumatorArial", "ResumatorArialBold"
    except Exception:
        pass
    return "Helvetica", "Helvetica-Bold"


def _prepare_lines(response_text: str, prompt_name: str | None, source_pdf: Path | Iterable[Path] | None) -> list[str]:
    header = [
        DOCUMENT_TITLE,
        *_metadata_lines(prompt_name, source_pdf),
    ]
    header.append("")

    text = response_text.strip() or "Sem resposta informada."
    body: list[str] = []
    for paragraph in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        if not paragraph.strip():
            body.append("")
            continue
        body.extend(
            textwrap.wrap(
                paragraph,
                width=95,
                break_long_words=True,
                replace_whitespace=False,
            )
        )
    return header + body


def _metadata_lines(prompt_name: str | None, source_pdf: Path | Iterable[Path] | None) -> list[str]:
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
    header = [
        f"Aplicativo: {APP_NAME}",
        f"Desenvolvedor: {DEVELOPER}",
        f"Gerado em: {generated_at}",
    ]
    if prompt_name:
        header.append(f"Prompt: {prompt_name}")
    source_pdfs = _normalize_source_pdfs(source_pdf)
    if len(source_pdfs) == 1:
        header.append(f"Arquivo analisado: {source_pdfs[0].name}")
    elif source_pdfs:
        header.append(f"Arquivos analisados: {', '.join(path.name for path in source_pdfs)}")
    return header


def _normalize_source_pdfs(source_pdf: Path | Iterable[Path] | None) -> list[Path]:
    if source_pdf is None:
        return []
    if isinstance(source_pdf, Path):
        return [source_pdf]
    return [Path(path) for path in source_pdf]


def _paginate(lines: list[str]) -> list[list[str]]:
    max_lines = int((TOP - BOTTOM) / LINE_HEIGHT) - 2
    pages: list[list[str]] = []
    current: list[str] = []

    for line in lines:
        if len(current) >= max_lines:
            pages.append(current)
            current = []
        current.append(line)

    pages.append(current or [""])
    return pages


def _build_pdf(pages: list[list[str]]) -> bytes:
    page_streams = [_page_content(lines, index + 1, len(pages)) for index, lines in enumerate(pages)]
    page_count = len(page_streams)

    font_regular_id = 3
    font_bold_id = 4
    first_content_id = 5
    first_page_id = first_content_id + page_count
    max_id = first_page_id + page_count - 1

    objects: dict[int, bytes] = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        3: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>",
        4: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold /Encoding /WinAnsiEncoding >>",
    }

    kids = []
    for index, stream in enumerate(page_streams):
        content_id = first_content_id + index
        page_id = first_page_id + index
        kids.append(f"{page_id} 0 R")
        objects[content_id] = (
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"\nendstream"
        )
        objects[page_id] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {PAGE_WIDTH} {PAGE_HEIGHT}] "
            f"/Resources << /Font << /F1 {font_regular_id} 0 R /F2 {font_bold_id} 0 R >> >> "
            f"/Contents {content_id} 0 R >>"
        ).encode("ascii")

    objects[2] = f"<< /Type /Pages /Kids [{' '.join(kids)}] /Count {page_count} >>".encode("ascii")

    output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0] * (max_id + 1)
    for obj_id in range(1, max_id + 1):
        offsets[obj_id] = len(output)
        output.extend(f"{obj_id} 0 obj\n".encode("ascii"))
        output.extend(objects[obj_id])
        output.extend(b"\nendobj\n")

    xref_at = len(output)
    output.extend(f"xref\n0 {max_id + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for obj_id in range(1, max_id + 1):
        output.extend(f"{offsets[obj_id]:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer\n<< /Size {max_id + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF\n".encode(
            "ascii"
        )
    )
    return bytes(output)


def _page_content(lines: list[str], page_number: int, total_pages: int) -> bytes:
    commands: list[bytes] = []
    commands.append(_draw_text(LEFT, 762, "F2", 13, APP_NAME))
    commands.append(_draw_text(LEFT, 744, "F1", 9, f"Página {page_number} de {total_pages}"))

    y = TOP - 28
    for line in lines:
        if line == DOCUMENT_TITLE:
            commands.append(_draw_text(LEFT, y, "F2", 13, line))
        else:
            commands.append(_draw_text(LEFT, y, "F1", BODY_FONT_SIZE, line))
        y -= LINE_HEIGHT

    return b"".join(commands)


def _draw_text(x: int, y: int, font: str, size: int, text: str) -> bytes:
    return (
        f"BT /{font} {size} Tf {x} {y} Td ".encode("ascii")
        + _pdf_string(text)
        + b" Tj ET\n"
    )


def _pdf_string(text: str) -> bytes:
    raw = text.encode("cp1252", errors="replace")
    raw = raw.replace(b"\\", b"\\\\").replace(b"(", b"\\(").replace(b")", b"\\)")
    return b"(" + raw + b")"
